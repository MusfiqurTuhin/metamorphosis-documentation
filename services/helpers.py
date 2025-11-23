import streamlit as st
import io
import json
import base64
import zlib
import requests
import re
from datetime import datetime
from fpdf import FPDF
from docx import Document
from PIL import Image
import os

def add_to_history(st_obj, feature, content, title="Untitled"):
    """Add item to history in session state."""
    if 'history' not in st_obj.session_state:
        st_obj.session_state.history = {}
    
    if feature not in st_obj.session_state.history:
        st_obj.session_state.history[feature] = []
    
    st_obj.session_state.history[feature].insert(0, {
        'title': title,
        'content': content,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
    })
    
    # Keep only last 10 items
    st_obj.session_state.history[feature] = st_obj.session_state.history[feature][:10]

def save_to_favorites(st_obj, content, feature, title):
    """Save item to favorites in session state."""
    if 'favorites' not in st_obj.session_state:
        st_obj.session_state.favorites = []
        
    st_obj.session_state.favorites.append({
        'feature': feature,
        'title': title,
        'content': content,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
    })

def extract_context_text(uploaded_file):
    """Extract text from uploaded file (txt, md, pdf, docx)."""
    if not uploaded_file:
        return None
    
    try:
        if uploaded_file.type == "text/plain" or uploaded_file.name.endswith(".md"):
            return uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"Error reading file: {str(e)}"
    return None

def create_docx(text):
    """Create a DOCX file from text."""
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf(text, image_bytes=None):
    """Create a PDF file from text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add image if provided
    if image_bytes:
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        try:
            pdf.image(tmp_path, x=10, w=190)
            pdf.ln(10)
        except:
            pdf.cell(0, 10, "Image Error", ln=True)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    # Add text
    # Encode/decode to handle latin-1 limitations of FPDF
    clean_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    return pdf.output(dest='S').encode('latin-1')

def sanitize_mermaid_code(raw_text):
    """Extract and sanitize Mermaid code from LLM response."""
    match = re.search(r"```mermaid\s+(.*?)\s+```", raw_text, re.DOTALL)
    if match:
        code = match.group(1).strip()
    else:
        code = raw_text.replace("```mermaid", "").replace("```", "").strip()
    return code

def get_mermaid_img(code, format="png"):
    """Generate Mermaid image using mermaid.ink API."""
    state = {
        "code": code,
        "mermaid": {"theme": "default", "securityLevel": "loose"},
        "autoSync": True,
        "updateDiagram": True
    }
    json_str = json.dumps(state)
    compressor = zlib.compressobj(9, zlib.DEFLATED, -15, 8, zlib.Z_DEFAULT_STRATEGY)
    compressed = compressor.compress(json_str.encode('utf-8')) + compressor.flush()
    base64_str = base64.urlsafe_b64encode(compressed).decode('utf-8')
    
    url = f"https://mermaid.ink/img/pako:{base64_str}"
    if format == "svg":
        url = f"https://mermaid.ink/svg/pako:{base64_str}"
    
    try:
        response = requests.get(url)
        if response.status_code == 200:
            return response.content
    except:
        return None
    return None

def convert_to_jpg(image_bytes):
    """Convert image bytes to JPG."""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        rgb_im = image.convert('RGB')
        output = io.BytesIO()
        rgb_im.save(output, format='JPEG', quality=95)
        return output.getvalue()
    except:
        return None
