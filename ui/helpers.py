import streamlit as st
import io
import json
import base64
import zlib
import requests
import re
from datetime import datetime
from fpdf import FPDF
from docx import Document
from PIL import Image
import os

def add_to_history(st_obj, feature, content, title="Untitled"):
    """Add item to history in session state."""
    if 'history' not in st_obj.session_state:
        st_obj.session_state.history = {}
    
    if feature not in st_obj.session_state.history:
        st_obj.session_state.history[feature] = []
    
    st_obj.session_state.history[feature].insert(0, {
        'title': title,
        'content': content,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
    })
    
    # Keep only last 10 items
    st_obj.session_state.history[feature] = st_obj.session_state.history[feature][:10]

def save_to_favorites(st_obj, content, feature, title):
    """Save item to favorites in session state."""
    if 'favorites' not in st_obj.session_state:
        st_obj.session_state.favorites = []
        
    st_obj.session_state.favorites.append({
        'feature': feature,
        'title': title,
        'content': content,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
    })

def extract_context_text(uploaded_file):
    """Extract text from uploaded file (txt, md, pdf, docx)."""
    if not uploaded_file:
        return None
    
    try:
        if uploaded_file.type == "text/plain" or uploaded_file.name.endswith(".md"):
            return uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"Error reading file: {str(e)}"
    return None

def create_docx(text):
    """Create a DOCX file from text."""
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf(text, image_bytes=None):
    """Create a PDF file from text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add image if provided
    if image_bytes:
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        try:
            pdf.image(tmp_path, x=10, w=190)
            pdf.ln(10)
        except:
            pdf.cell(0, 10, "Image Error", ln=True)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    # Add text
    # Encode/decode to handle latin-1 limitations of FPDF
    clean_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    return pdf.output(dest='S').encode('latin-1')

def sanitize_mermaid_code(raw_text):
    """Extract and sanitize Mermaid code from LLM response."""
    match = re.search(r"```mermaid\s+(.*?)\s+```", raw_text, re.DOTALL)
    if match:
        code = match.group(1).strip()
    else:
        code = raw_text.replace("```mermaid", "").replace("```", "").strip()
    return code

def get_kroki_img(code, format="png"):
    """
    Generate diagram using Kroki.io API.
    Uses POST request to avoid URL length limits for large diagrams.
    """
    try:
        url = f"https://kroki.io/mermaid/{format}"
        # Kroki POST expects raw diagram source in body (uncompressed is fine and standard for POST)
        response = requests.post(url, data=code.encode('utf-8'))
        
        if response.status_code == 200:
            return response.content
    except:
        return None
    return None

def get_mermaid_img(code, format="png", theme="default"):
    """
    Generate Mermaid image.
    Attempts Kroki.io first (more robust), falls back to mermaid.ink.
    """
    # 1. Try Kroki (Primary)
    # Note: Kroki doesn't support 'theme' via URL easily for Mermaid, it renders default.
    # But it handles complex syntax much better.
    # We inject theme directive into the code if possible
    kroki_code = code
    if "%%{init:" not in code and theme != "default":
        # Inject theme directive
        theme_json = json.dumps({"theme": theme})
        kroki_code = f"%%{{init: {theme_json} }}%%\n{code}"
        
    img = get_kroki_img(kroki_code, format)
    if img:
        return img
        
    # 2. Fallback to mermaid.ink
    state = {
        "code": code,
        "mermaid": {"theme": theme, "securityLevel": "loose"},
        "autoSync": True,
        "updateDiagram": True
    }
    json_str = json.dumps(state)
    compressor = zlib.compressobj(9, zlib.DEFLATED, -15, 8, zlib.Z_DEFAULT_STRATEGY)
    compressed = compressor.compress(json_str.encode('utf-8')) + compressor.flush()
    # Strip padding from base64 string, as required by mermaid.ink/pako
    base64_str = base64.urlsafe_b64encode(compressed).decode('utf-8').rstrip('=')
    
    url = f"https://mermaid.ink/img/pako:{base64_str}"
    if format == "svg":
        url = f"https://mermaid.ink/svg/pako:{base64_str}"
    
    try:
        response = requests.get(url)
        if response.status_code == 200:
            return response.content
    except:
        return None
    return None

def convert_to_jpg(image_bytes):
    """Convert image bytes to JPG."""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        rgb_im = image.convert('RGB')
        output = io.BytesIO()
        rgb_im.save(output, format='JPEG', quality=95)
        return output.getvalue()
    except:
        return None

def fix_mermaid_syntax(code):
    """
    Apply regex-based auto-fixes to common Mermaid errors.
    This runs BEFORE the LLM validation loop as a fast correction layer.
    """
    fixed_lines = []
    lines = code.split('\n')
    
    is_mindmap = "mindmap" in code.lower()
    is_flowchart = "flowchart" in code.lower() or "graph" in code.lower()
    
    for line in lines:
        stripped = line.strip()
        
        # Global Fix: Strip inline comments (%%) from end of lines
        # But be careful not to strip %% inside quotes? 
        # Simpler approach: If line contains `%%`, remove everything after it, UNLESS the line is just a full comment
        if "%%" in line and not stripped.startswith("%%"):
             # Split and keep first part
             parts = line.split("%%")
             # Retain original indentation of the left part
             line = parts[0].rstrip() 
             
        # Fix 1: Mindmap node with nested parens/commas that are NOT quoted.
        if is_mindmap:
            # Handle Double Parentheses specially: root((Text))
            # 1. Check for double parens explicitly first
            # regex: (prefix)((content))(suffix)
            double_match = re.search(r'^(\s*[^(\n]*)\(\(([^"]*)\)\)(\s*)$', line)
            if double_match:
                 # Check if content needs quoting (has special chars or is just plain)
                 # Actually, for consistency, we can just quote everything inside ((...)) if it's not already quoted?
                 # But usually ((Circle)) is fine. The issue is if content has special chars.
                 # Let's simple-check: IF content has ( or ) or , AND is not quoted -> Quote it.
                 content = double_match.group(2)
                 if any(c in content for c in "(),") and not (content.startswith('"') and content.endswith('"')):
                     prefix = double_match.group(1)
                     suffix = double_match.group(3)
                     fixed_line = f'{prefix}(("{content}")){suffix}'
                     fixed_lines.append(fixed_line)
                     continue

            # 2. Standard Mindmap: ID(Text with (parens) or , commas) -> ID("Text ...")
            # ALLOW spaces in ID precursor: ^(\s*[^(\n]+)\(
            # This captures: "  My Node ID (Content)"
            match = re.search(r'^(\s*[^(\n]*)\(([^"]*[\(\),][^"]*)\)(\s*)$', line)
            if match:
                # We need to make sure this wasn't actually a double paren case that slipped through?
                # The previous regex `[^"]*` is greedy. 
                # If the line is `root((Garment))`, match group 1 is `root(`, group 2 is `Garment`.
                # This logic BREAKS valid double parens if we blindly quote.
                # FIX: Verify that the char BEFORE `(` is NOT `(` and char AFTER `)` is NOT `)`?
                # Easier: Check if line contains `((` and skip this single-paren Fix?
                if "((" in line:
                    fixed_lines.append(line)
                    continue

                prefix = match.group(1)
                content = match.group(2)
                suffix = match.group(3)
                # Apply fix
                fixed_line = f'{prefix}("{content}"){suffix}'
                fixed_lines.append(fixed_line)
                continue
                
        # Fix 2: Flowchart node label quoting
        if is_flowchart:
            # Pattern: ID[Text with (parens) or special chars] -> ID["Text ..."]
            # Regex: (ID) [ (Content) ] (Suffix)
            # Find unquoted content inside [] that contains ( or )
            match = re.search(r'(\S+)\[([^"\]]*[\(\)][^"\]]*)\]', line)
            if match:
                # Reconstruct line with quotes around the content inside []
                # Note: This is a robust replacement for the match
                start, end = match.span()
                node_id = match.group(1)
                content = match.group(2)
                
                # Check strictness: if content is already quoted? Regex `[^"]` handles it partially.
                # Replace the match segment
                new_segment = f'{node_id}["{content}"]'
                line = line[:start] + new_segment + line[end:]
            
        fixed_lines.append(line)
        
    return '\n'.join(fixed_lines)

def validate_mermaid_syntax(code):
    """
    Validate Mermaid syntax and return a list of specific errors/warnings.
    Based on common LLM failure modes.
    """
    errors = []
    
    # 1. Check for spaces after commas in style/classDefs (e.g. "fill:#f9f, stroke:#333")
    if re.search(r":[#a-zA-Z0-9]+,\s+", code):
        errors.append("❌ Style Error: Remove spaces after commas in style definitions (e.g., use 'fill:#fff,stroke:#000', NOT 'fill:#fff, stroke:#000').")
        
    # 2. Check for unclosed blocks
    # Keywords that require an 'end'
    block_keywords = ['subgraph', 'loop', 'opt', 'alt', 'par', 'rect', 'critical', 'break', 'parallel']
    
    total_opens = 0
    total_ends = len(re.findall(r'\\bend\\b', code, re.IGNORECASE))
    
    for kw in block_keywords:
        opens = len(re.findall(f'\\b{kw}\\b', code, re.IGNORECASE))
        total_opens += opens
        
    if total_opens != total_ends:
        errors.append(f"❌ Block Error: Found {total_opens} opening blocks but {total_ends} 'end' statements. Check if all subgraphs/loops are closed.")

    # 3. Gantt Layout Specifics
    if "gantt" in code.lower() and re.search(r'^\s*today\b', code, re.MULTILINE | re.IGNORECASE):
        errors.append("❌ Gantt Error: Found line starting with 'today'. Mermaid does NOT support defining 'today' manually using a date. REMOVE this line completely.")
    
    # Check for invalid comments (//)
    if "//" in code:
        errors.append("❌ Syntax Error: Mermaid uses '%%' for comments, not '//'. Replace '//' with '%%' or remove the comment.")

    # 4. Mindmap specific: Check for text after node definition on the same line
    if "mindmap" in code.lower():
         lines = code.split('\n')
         for i, line in enumerate(lines):
             # Match lines that look like:  ID(Text)AnyThingElse
             # Updated regex to be ultra strict:
             # Look for a closing valid bracket [ ) ] } that terminates the node definition.
             # Then check if there is ANY non-whitespace character after it on the same line.
             # This catches: Node(Text)Text, Node(Text) Text, Node("Text")Text
             if re.search(r'[\]\)\}][\t ]*\S', line):
                 # Filter out false positives like "Node(A) --> Node(B)" which is valid in some diagrams but not mindmap text based
                 if "-->" not in line and "---" not in line:
                    errors.append(f"❌ Mindmap Error (Line {i+1}): Found text after node definition. Ensure each node is on its own line. (Content: '{line.strip()}')")
                    break
             
             # Check for unquoted special characters inside the node text (e.g. Node(Text (Detail)))
             # This creates a basic check: if we see '(', ')', or ',' inside the node content AND it's not wrapped in quotes.
             # Regex detects: ID( ... [(),] ... ) where no quotes " are present.
             # We look for: non-quote chars, then special char, then non-quote chars inside brackets
             if re.search(r'[\(\[\{][^"]*[\(\),][^"]*[\)\]\}]', line):
                 # Refined check: does the line contain double quotes?
                 if '"' not in line:
                     errors.append(f"❌ Mindmap Error (Line {i+1}): Text containing brackets '()' or commas MUST be wrapped in double quotes. (e.g. use `Node(\"Text (Detail)\")` instead of `Node(Text (Detail))`)")
                     break # Trigger fix

    # 5. Gantt Specific: Check for invalid date formats and line structure
    if "gantt" in code.lower():
        lines = code.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or line.startswith("title") or line.startswith("dateFormat") or line.startswith("axisFormat") or line.startswith("section") or line.startswith("%%"):
                continue
            
            # Simple heuristic for tasks: "Task Name : [crit, active, done etc], [after x], 2023-01-01, 30d"
            # We want to catch lines that don't look like sections or metadata and are missing dates or duration
            if ":" in line:
                 # Check for YYYY-MM-DD
                 if not re.search(r'\d{4}-\d{2}-\d{2}', line) and not re.search(r'\d+[dwms]', line) and "after" not in line:
                      errors.append(f"❌ Gantt Error (Line {i+1}): Task seems to be missing a start date (YYYY-MM-DD) or duration (e.g. 5d).")

    # 6. Flowchart Specific: Check for broken arrows
    if "flowchart" in code.lower() or "graph" in code.lower():
         lines = code.split('\n')
         for i, line in enumerate(lines):
             line = line.strip()
             # Check if line ends with an arrow operator without a target
             # Operators: -->, ---, -.->, ==>
             if re.search(r'(-{2,}>?|\.-+>|={2,}>)\s*$', line):
                 errors.append(f"❌ Flowchart Error (Line {i+1}): Line ends with a connector/arrow but has no target node. (Content: '{line}')")

    # 7. ER Diagram Specific: Check for valid attribute definitions
    if "erDiagram" in code.lower():
        lines = code.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            # Skip relationships (contains --) and entity block starts ({)
            if regex_relationship := re.search(r'[\}\|o]--[\}\|o]', line): continue
            if "{" in line or "}" in line: continue
            
            # Attribute line usually: type name [PK|FK] [comment]
            # Bad pattern seen: "Packing_WCID FK VARCHAR S" -> Name Key Type?
            # We want to enforce: Type Name [Key]
            # Let's just catch lines that have many words but don't look right
            words = line.split()
            if len(words) >= 3:
                # If the second word is FK or PK, it might be flipped. 
                # Standard is: string name PK
                # Bad: name PK string
                if words[1] in ["PK", "FK"] and len(words) > 2:
                     # heuristic: formatting error
                     errors.append(f"❌ ER Diagram Error (Line {i+1}): Attribute seems malformed. Format should be `Type Name [PK/FK]`. (Found: '{line}')")

    return errors
